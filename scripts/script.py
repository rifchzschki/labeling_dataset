import random
from openpyxl import load_workbook
import docx
import pdfkit
from PIL import Image
from spire.doc import *
from spire.doc.common import *

def extract_replacements(excel_path):
    wb = load_workbook(excel_path)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    headers = rows[0]
    random_row = random.choice(rows[1:])
    replacements = {headers[i]: random_row[i] for i in range(len(headers))}
    
    return replacements

def replace_text_in_docx(docx_path, replacements, output_path):
    doc = docx.Document(docx_path)

    for para in doc.paragraphs:
        for key, value in replacements.items():
            if f'<{key}>' in para.text:
                para.text = para.text.replace(f'<{key}>', str(value))
                print(f'<{key}> ' + str(value))

    doc.save(output_path)

def generate_random_coordinate():
    latitude = random.uniform(-10.000000, 10.000000)
    longitude = random.uniform(110.000000, 120.000000)
    return f"{latitude:.6f}, {longitude:.6f}"

def convert_docx_to_pdf(docx_path, pdf_path):
    options = {
        'quiet': '',
    }
    pdfkit.from_file(docx_path, pdf_path, options=options)

def convert_pdf_to_jpeg(pdf_path, jpeg_path):
    images = Image.open(pdf_path)
    images.convert('RGB').save(jpeg_path, 'JPEG')

def main():
    excel_path = './slo-excel-data.xlsx'
    input_docx_path = './slo-template-edited.docx'

    os.makedirs('./output_docx', exist_ok=True)
    os.makedirs('./inputs_revised', exist_ok=True)

    # replacements = extract_replacements(excel_path)

    # replacements['koordinat'] = generate_random_coordinate()
    # replacements['titik'] = random.randint(1, 5)

    # print('here2')
    # output_docx_path = 'output.docx'

    # replace_text_in_docx(input_docx_path, replacements, output_docx_path)

    # print(f"{output_docx_path}")

    for i in range(1, 101):
        replacements = extract_replacements(excel_path)

        replacements['koordinat'] = generate_random_coordinate()

        replacements['titik_1'] = random.randint(1, 5)
        replacements['titik_2'] = random.randint(1, 5)
        replacements['titik_3'] = random.randint(1, 5)
        replacements['titik_4'] = random.randint(1, 5)
        replacements['titik_5'] = random.randint(1, 5)
        replacements['titik_6'] = random.randint(1, 5)

        if len(replacements['alamat_instalasi'])>8:
            replacements['alamat_instalasi'] = replacements['alamat_instalasi'][:8]

        if len(replacements['nama_provinsi'])>8:
            replacements['nama_provinsi'] = replacements['nama_provinsi'][:8]

        output_docx_path = f'./output_docx/output_{i}.docx'
        replace_text_in_docx(input_docx_path, replacements, output_docx_path)

        # output_pdf_path = f'output_{i}.pdf'
        # convert_docx_to_pdf(output_docx_path, output_pdf_path)

        # output_jpeg_path = f'./inputs/input_{i}.jpeg'
        # convert_pdf_to_jpeg(output_pdf_path, output_jpeg_path)

        # import aspose.words as aw

        # doc = aw.Document(output_docx_path)

        # for page in range(0, doc.page_count):
        #     extractedPage = doc.extract_pages(page, 1)
        #     extractedPage.save(f"./inputs/input_{page + 1}.jpg")

        # document = Document()
        # document.LoadFromFile(output_docx_path)

        # # Convert the document to a list of image streams
        # image_streams = document.SaveImageToStreams(ImageType.Bitmap)

        # # Save each image stream to a PNG file
        # for image in image_streams:
        #     image_name = f"./inputs/input_{i}.png"
        #     with open(image_name,'wb') as image_file:
        #         image_file.write(image.ToArray())

        # # Close the document
        # document.Close()

if __name__ == "__main__":
    main()